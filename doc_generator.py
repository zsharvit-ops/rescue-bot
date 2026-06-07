"""
Generates a .docx search & rescue report with:
  Summary — operational quick view
  Part 1 — structured table of extracted fields
  Part 2 — full original transcript
"""

import os
import zipfile
import shutil
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from extractor import FIELDS


def _force_rtl_in_zip(docx_path: str):
    """Post-process the docx zip to inject RTL into every paragraph in document.xml."""
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    xml = data.decode("utf-8")
                    # Add <w:bidi/> inside <w:pPr> only if not already there
                    def add_bidi(m):
                        inner = m.group(0)
                        if "w:bidi" not in inner:
                            inner = inner.replace("</w:pPr>", "<w:bidi/><w:jc w:val=\"right\"/></w:pPr>")
                        return inner
                    xml = re.sub(r"<w:pPr>.*?</w:pPr>", add_bidi, xml, flags=re.DOTALL)
                    # Add <w:rtl/> inside <w:rPr> only if not already there
                    def add_rtl(m):
                        inner = m.group(0)
                        if "w:rtl" not in inner:
                            inner = inner.replace("</w:rPr>", "<w:rtl/></w:rPr>")
                        return inner
                    xml = re.sub(r"<w:rPr>.*?</w:rPr>", add_rtl, xml, flags=re.DOTALL)
                    data = xml.encode("utf-8")
                elif item.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "<w:bidi" not in xml:
                        xml = xml.replace("</w:settings>", "<w:bidi/></w:settings>")
                    data = xml.encode("utf-8")
                elif item.filename == "word/styles.xml":
                    xml = data.decode("utf-8")
                    # Force RTL on every <w:pPr> block inside styles too
                    def add_bidi_style(m):
                        inner = m.group(0)
                        if "w:bidi" not in inner:
                            inner = inner.replace("</w:pPr>", '<w:bidi/><w:jc w:val="right"/></w:pPr>')
                        return inner
                    xml = re.sub(r"<w:pPr>.*?</w:pPr>", add_bidi_style, xml, flags=re.DOTALL)
                    # Also add RTL lang to every <w:rPr> in styles
                    def add_rtl_style(m):
                        inner = m.group(0)
                        if "w:rtl" not in inner:
                            inner = inner.replace("</w:rPr>", '<w:rtl/><w:cs/></w:rPr>')
                        return inner
                    xml = re.sub(r"<w:rPr>.*?</w:rPr>", add_rtl_style, xml, flags=re.DOTALL)
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
    os.replace(tmp, docx_path)


def _rtl_paragraph(paragraph):
    """Apply RTL + right-align to a paragraph (call AFTER adding text)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    for tag in (qn("w:bidi"), qn("w:jc")):
        for el in pPr.findall(tag):
            pPr.remove(el)
    pPr.append(OxmlElement("w:bidi"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        for tag in (qn("w:rtl"), qn("w:lang")):
            for el in rPr.findall(tag):
                rPr.remove(el)
        rPr.append(OxmlElement("w:rtl"))
        lang = OxmlElement("w:lang")
        lang.set(qn("w:bidi"), "he-IL")
        rPr.append(lang)


def _rtl_cell(cell):
    for p in cell.paragraphs:
        _rtl_paragraph(p)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_rtl(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    _rtl_paragraph(p)
    return p


def generate_doc(fields: dict, transcript: str, summary: dict = None) -> str:
    doc = Document()
    doc.sections[0].page_width = Cm(21)
    doc.sections[0].page_height = Cm(29.7)

    try:
        normal = doc.styles["Normal"]
        pPr = normal.element.get_or_add_pPr()
        for tag in (qn("w:bidi"), qn("w:jc")):
            for el in pPr.findall(tag):
                pPr.remove(el)
        pPr.append(OxmlElement("w:bidi"))
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)
    except Exception:
        pass

    title_p = doc.add_paragraph()
    run = title_p.add_run("דוח תחקיר חילוץ והצלה")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    _rtl_paragraph(title_p)

    date_p = doc.add_paragraph()
    date_p.add_run(f"תאריך הפקה: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    _rtl_paragraph(date_p)

    doc.add_paragraph()

    if summary:
        _add_heading(doc, "סיכום מצב — מידע קריטי", level=1)
        sum_table = doc.add_table(rows=0, cols=2)
        sum_table.style = "Table Grid"
        sum_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        _set_table_rtl(sum_table)
        summary_fields = [
            ("תיאור המקרה",      summary.get("case_summary", "")),
            ("מיקום אחרון ידוע", summary.get("last_location", "")),
            ("מצב רפואי",        summary.get("medical_status", "")),
            ("זמן מאז נעדר",     summary.get("time_missing", "")),
            ("תיאור אישי",       summary.get("physical_description", "")),
        ]
        for label, value in summary_fields:
            if value:
                row = sum_table.add_row().cells
                row[0].text = str(value)
                row[1].text = label
                _set_cell_bg(row[1], "1A5376")
                for p in row[1].paragraphs:
                    _rtl_paragraph(p)
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(11)
                _set_cell_bg(row[0], "EBF5FB")
                for p in row[0].paragraphs:
                    _rtl_paragraph(p)
                    for r in p.runs:
                        r.font.size = Pt(11)
        for row in sum_table.rows:
            row.cells[0].width = Cm(12)
            row.cells[1].width = Cm(5)
        doc.add_paragraph()

    _add_heading(doc, "חלק א׳ — פרטי מחולץ ומקרה", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_rtl(table)

    hdr = table.rows[0].cells
    hdr[0].text = "ערך"
    hdr[1].text = "שדה"
    for cell in hdr:
        _set_cell_bg(cell, "1A5376")
        for p in cell.paragraphs:
            _rtl_paragraph(p)
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(11)

    for idx, (hebrew_label, key) in enumerate(FIELDS):
        value = str(fields.get(key) or "")
        row = table.add_row().cells
        row[0].text = value
        row[1].text = hebrew_label
        bg = "D6EAF8" if idx % 2 == 0 else "FFFFFF"
        for cell in row:
            _set_cell_bg(cell, bg)
            _rtl_cell(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    for row in table.rows:
        row.cells[0].width = Cm(11)
        row.cells[1].width = Cm(6)

    doc.add_page_break()

    _add_heading(doc, "חלק ב׳ — תמלול השיחה המלא", level=1)
    for line in transcript.splitlines():
        p = doc.add_paragraph(line or " ")
        _rtl_paragraph(p)
        for r in p.runs:
            r.font.size = Pt(10)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    os.makedirs("output", exist_ok=True)
    out_path = os.path.join("output", f"rescue_report_{timestamp}.docx")
    doc.save(out_path)
    _force_rtl_in_zip(out_path)
    return out_path
